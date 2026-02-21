import io
import tempfile
import zipfile
from typing import List, Dict, Any, Optional, Tuple, Union
from pathlib import Path
import PyPDF2
from docx import Document
from docx.shared import Inches
import logging
import subprocess
import os
import json
from PIL import Image

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Advanced document processing service for DOCX, XLSX, and PDF conversions"""
    
    @staticmethod
    async def pdf_to_docx(
        pdf_data: bytes,
        preserve_formatting: bool = True,
        extract_images: bool = False
    ) -> bytes:
        """Convert PDF to DOCX format"""
        try:
            # Extract text from PDF
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(pdf_data))
            
            # Create new DOCX document
            doc = Document()
            
            # Add document title
            title = doc.add_heading('Converted from PDF', 0)
            
            # Process each page
            for page_num, page in enumerate(pdf_reader.pages, 1):
                try:
                    # Extract text from page
                    page_text = page.extract_text()
                    
                    if page_text.strip():
                        # Add page heading
                        if len(pdf_reader.pages) > 1:
                            doc.add_heading(f'Page {page_num}', level=1)
                        
                        # Split text into paragraphs
                        paragraphs = page_text.split('\n\n')
                        
                        for paragraph_text in paragraphs:
                            if paragraph_text.strip():
                                # Clean up text
                                cleaned_text = ' '.join(paragraph_text.split())
                                if cleaned_text:
                                    doc.add_paragraph(cleaned_text)
                    
                    # Add page break between pages (except last page)
                    if page_num < len(pdf_reader.pages):
                        doc.add_page_break()
                        
                except Exception as e:
                    logger.warning(f"Failed to process page {page_num}: {e}")
                    doc.add_paragraph(f"[Error processing page {page_num}]")
            
            # Save document to bytes
            doc_buffer = io.BytesIO()
            doc.save(doc_buffer)
            doc_buffer.seek(0)
            
            logger.info(f"Converted PDF to DOCX with {len(pdf_reader.pages)} pages")
            return doc_buffer.getvalue()
            
        except Exception as e:
            logger.error(f"PDF to DOCX conversion failed: {e}")
            raise
    
    @staticmethod
    async def docx_to_pdf(docx_data: bytes) -> bytes:
        """Convert DOCX to PDF using LibreOffice"""
        try:
            with tempfile.TemporaryDirectory() as temp_dir:
                # Save DOCX to temporary file
                docx_path = os.path.join(temp_dir, "input.docx")
                with open(docx_path, 'wb') as f:
                    f.write(docx_data)
                
                # Convert using LibreOffice headless
                cmd = [
                    'libreoffice',
                    '--headless',
                    '--convert-to', 'pdf',
                    '--outdir', temp_dir,
                    docx_path
                ]
                
                result = subprocess.run(cmd, capture_output=True, text=True, timeout=60)
                
                if result.returncode != 0:
                    raise Exception(f"LibreOffice conversion failed: {result.stderr}")
                
                # Read converted PDF
                pdf_path = os.path.join(temp_dir, "input.pdf")
                if not os.path.exists(pdf_path):
                    raise Exception("PDF output file not created")
                
                with open(pdf_path, 'rb') as f:
                    pdf_data = f.read()
                
                logger.info("Converted DOCX to PDF successfully")
                return pdf_data
                
        except subprocess.TimeoutExpired:
            raise Exception("DOCX to PDF conversion timed out")
        except Exception as e:
            logger.error(f"DOCX to PDF conversion failed: {e}")
            raise
    
    @staticmethod
    async def extract_docx_content(docx_data: bytes) -> Dict[str, Any]:
        """Extract content and metadata from DOCX"""
        try:
            doc = Document(io.BytesIO(docx_data))
            
            # Extract text content
            paragraphs = []
            tables = []
            
            for element in doc.element.body:
                if element.tag.endswith('p'):  # Paragraph
                    para = None
                    for p in doc.paragraphs:
                        if p._element == element:
                            para = p
                            break
                    if para and para.text.strip():
                        paragraphs.append({
                            'text': para.text,
                            'style': para.style.name if para.style else 'Normal'
                        })
                
                elif element.tag.endswith('tbl'):  # Table
                    table = None
                    for t in doc.tables:
                        if t._element == element:
                            table = t
                            break
                    if table:
                        table_data = []
                        for row in table.rows:
                            row_data = [cell.text.strip() for cell in row.cells]
                            table_data.append(row_data)
                        tables.append(table_data)
            
            # Extract metadata
            core_props = doc.core_properties
            metadata = {
                'title': core_props.title or '',
                'author': core_props.author or '',
                'subject': core_props.subject or '',
                'created': str(core_props.created) if core_props.created else '',
                'modified': str(core_props.modified) if core_props.modified else '',
                'revision': core_props.revision or 0
            }
            
            # Count statistics
            full_text = ' '.join([p['text'] for p in paragraphs])
            word_count = len(full_text.split())
            char_count = len(full_text)
            
            return {
                'paragraphs': paragraphs,
                'tables': tables,
                'metadata': metadata,
                'statistics': {
                    'paragraph_count': len(paragraphs),
                    'table_count': len(tables),
                    'word_count': word_count,
                    'character_count': char_count
                },
                'full_text': full_text
            }
            
        except Exception as e:
            logger.error(f"DOCX content extraction failed: {e}")
            raise
    
    @staticmethod
    async def create_xlsx_from_data(data: List[List[str]], sheet_name: str = "Sheet1") -> bytes:
        """Create XLSX file from tabular data using openpyxl"""
        try:
            from openpyxl import Workbook
            
            # Create workbook and worksheet
            wb = Workbook()
            ws = wb.active
            ws.title = sheet_name
            
            # Add data to worksheet
            for row_idx, row_data in enumerate(data, 1):
                for col_idx, cell_value in enumerate(row_data, 1):
                    ws.cell(row=row_idx, column=col_idx, value=str(cell_value))
            
            # Save to bytes
            xlsx_buffer = io.BytesIO()
            wb.save(xlsx_buffer)
            xlsx_buffer.seek(0)
            
            logger.info(f"Created XLSX with {len(data)} rows using openpyxl")
            return xlsx_buffer.getvalue()
                
        except Exception as e:
            logger.error(f"XLSX creation failed: {e}")
            raise
    
    @staticmethod
    async def xlsx_to_csv(xlsx_data: bytes) -> bytes:
        """Convert XLSX to CSV format using openpyxl"""
        try:
            from openpyxl import load_workbook
            import csv
            
            # Load workbook
            wb = load_workbook(io.BytesIO(xlsx_data))
            ws = wb.active
            
            # Convert to CSV
            csv_buffer = io.StringIO()
            writer = csv.writer(csv_buffer)
            
            for row in ws.iter_rows(values_only=True):
                # Convert None values to empty strings
                row_data = [str(cell) if cell is not None else '' for cell in row]
                writer.writerow(row_data)
            
            # Convert to bytes
            csv_content = csv_buffer.getvalue()
            csv_bytes = csv_content.encode('utf-8')
            
            logger.info("Converted XLSX to CSV successfully using openpyxl")
            return csv_bytes
                
        except Exception as e:
            logger.error(f"XLSX to CSV conversion failed: {e}")
            raise
    
    @staticmethod
    async def merge_documents(
        documents: List[Tuple[bytes, str]],  # (data, format)
        output_format: str = 'pdf'
    ) -> bytes:
        """Merge multiple documents into one"""
        try:
            if output_format == 'pdf':
                # Convert all documents to PDF first, then merge
                pdf_data_list = []
                
                for doc_data, doc_format in documents:
                    if doc_format.lower() == 'pdf':
                        pdf_data_list.append(doc_data)
                    elif doc_format.lower() == 'docx':
                        pdf_data = await DocumentProcessor.docx_to_pdf(doc_data)
                        pdf_data_list.append(pdf_data)
                    else:
                        raise Exception(f"Unsupported format for merging: {doc_format}")
                
                # Use existing PDF merge functionality
                from .pdf_processor import PDFProcessor
                merged_pdf = await PDFProcessor.merge_pdfs(pdf_data_list)
                
                logger.info(f"Merged {len(documents)} documents into PDF")
                return merged_pdf
                
            elif output_format == 'docx':
                # Merge into DOCX
                merged_doc = Document()
                merged_doc.add_heading('Merged Document', 0)
                
                for i, (doc_data, doc_format) in enumerate(documents):
                    if i > 0:
                        merged_doc.add_page_break()
                    
                    merged_doc.add_heading(f'Document {i + 1}', level=1)
                    
                    if doc_format.lower() == 'docx':
                        # Extract content from DOCX
                        content = await DocumentProcessor.extract_docx_content(doc_data)
                        for para in content['paragraphs']:
                            merged_doc.add_paragraph(para['text'])
                    elif doc_format.lower() == 'pdf':
                        # Convert PDF to text and add
                        from .pdf_processor import PDFProcessor
                        text_result = await PDFProcessor.extract_text(doc_data)
                        merged_doc.add_paragraph(text_result['text'])
                    else:
                        merged_doc.add_paragraph(f"[Unsupported format: {doc_format}]")
                
                # Save merged document
                doc_buffer = io.BytesIO()
                merged_doc.save(doc_buffer)
                doc_buffer.seek(0)
                
                logger.info(f"Merged {len(documents)} documents into DOCX")
                return doc_buffer.getvalue()
            
            else:
                raise Exception(f"Unsupported output format: {output_format}")
                
        except Exception as e:
            logger.error(f"Document merge failed: {e}")
            raise
    
    @staticmethod
    async def split_document(
        doc_data: bytes,
        doc_format: str,
        split_method: str = 'pages',
        split_params: Dict[str, Any] = None
    ) -> List[bytes]:
        """Split document based on various criteria"""
        try:
            if doc_format.lower() == 'pdf':
                # Use existing PDF split functionality
                from .pdf_processor import PDFProcessor
                
                if split_method == 'pages' and split_params:
                    page_ranges = split_params.get('page_ranges', [])
                    return await PDFProcessor.split_pdf(doc_data, page_ranges)
                else:
                    raise Exception("PDF splitting requires page ranges")
                    
            elif doc_format.lower() == 'docx':
                # Split DOCX by sections or pages
                doc = Document(io.BytesIO(doc_data))
                
                if split_method == 'sections':
                    # Split by headings
                    sections = []
                    current_section = Document()
                    current_heading = None
                    
                    for para in doc.paragraphs:
                        if para.style.name.startswith('Heading'):
                            if current_heading is not None:
                                # Save current section
                                section_buffer = io.BytesIO()
                                current_section.save(section_buffer)
                                section_buffer.seek(0)
                                sections.append(section_buffer.getvalue())
                                
                                # Start new section
                                current_section = Document()
                            
                            current_heading = para.text
                            current_section.add_heading(para.text, level=1)
                        else:
                            if para.text.strip():
                                current_section.add_paragraph(para.text)
                    
                    # Save last section
                    if current_heading is not None:
                        section_buffer = io.BytesIO()
                        current_section.save(section_buffer)
                        section_buffer.seek(0)
                        sections.append(section_buffer.getvalue())
                    
                    logger.info(f"Split DOCX into {len(sections)} sections")
                    return sections
                
                else:
                    raise Exception(f"Unsupported split method for DOCX: {split_method}")
            
            else:
                raise Exception(f"Unsupported format for splitting: {doc_format}")
                
        except Exception as e:
            logger.error(f"Document split failed: {e}")
            raise
    
    @staticmethod
    async def get_document_info(doc_data: bytes, doc_format: str) -> Dict[str, Any]:
        """Get comprehensive document information"""
        try:
            if doc_format.lower() == 'pdf':
                from .pdf_processor import PDFProcessor
                return await PDFProcessor.get_pdf_info(doc_data)
                
            elif doc_format.lower() == 'docx':
                content = await DocumentProcessor.extract_docx_content(doc_data)
                
                return {
                    'format': 'DOCX',
                    'file_size': len(doc_data),
                    'paragraph_count': content['statistics']['paragraph_count'],
                    'table_count': content['statistics']['table_count'],
                    'word_count': content['statistics']['word_count'],
                    'character_count': content['statistics']['character_count'],
                    'metadata': content['metadata']
                }
                
            elif doc_format.lower() == 'xlsx':
                from openpyxl import load_workbook
                
                wb = load_workbook(io.BytesIO(doc_data))
                ws = wb.active
                
                # Count rows and columns with data
                max_row = ws.max_row
                max_col = ws.max_column
                
                # Count non-empty cells
                non_empty_cells = 0
                for row in ws.iter_rows(values_only=True):
                    for cell in row:
                        if cell is not None and str(cell).strip():
                            non_empty_cells += 1
                
                return {
                    'format': 'XLSX',
                    'file_size': len(doc_data),
                    'type': 'spreadsheet',
                    'sheet_count': len(wb.sheetnames),
                    'sheet_names': wb.sheetnames,
                    'active_sheet': ws.title,
                    'max_row': max_row,
                    'max_column': max_col,
                    'non_empty_cells': non_empty_cells,
                    'metadata': {
                        'title': wb.properties.title or '',
                        'creator': wb.properties.creator or '',
                        'created': str(wb.properties.created) if wb.properties.created else '',
                        'modified': str(wb.properties.modified) if wb.properties.modified else ''
                    }
                }
            
            else:
                return {
                    'format': doc_format.upper(),
                    'file_size': len(doc_data),
                    'supported': False
                }
                
        except Exception as e:
            logger.error(f"Failed to get document info: {e}")
            raise
    
    @staticmethod
    async def convert_document_format(
        doc_data: bytes,
        input_format: str,
        output_format: str,
        options: Dict[str, Any] = None
    ) -> bytes:
        """Universal document format converter"""
        try:
            input_fmt = input_format.lower()
            output_fmt = output_format.lower()
            options = options or {}
            
            # PDF conversions
            if input_fmt == 'pdf' and output_fmt == 'docx':
                return await DocumentProcessor.pdf_to_docx(
                    doc_data,
                    preserve_formatting=options.get('preserve_formatting', True),
                    extract_images=options.get('extract_images', False)
                )
            
            # DOCX conversions
            elif input_fmt == 'docx' and output_fmt == 'pdf':
                return await DocumentProcessor.docx_to_pdf(doc_data)
            
            # XLSX conversions
            elif input_fmt == 'xlsx' and output_fmt == 'csv':
                return await DocumentProcessor.xlsx_to_csv(doc_data)
            
            # Same format (no conversion needed)
            elif input_fmt == output_fmt:
                return doc_data
            
            else:
                raise Exception(f"Conversion from {input_fmt} to {output_fmt} not supported")
                
        except Exception as e:
            logger.error(f"Document format conversion failed: {e}")
            raise